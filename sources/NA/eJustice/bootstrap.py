#!/usr/bin/env python3
"""
NA/eJustice -- Namibia eJustice Portal (Supreme Court + High Court Judgments)

Fetches Namibian court judgments from the official eJustice SharePoint portal
(https://ejustice.jud.na). Judgments are uploaded as Microsoft Word (.docx /
.doc) documents in folders by court and category. The folder browse URLs
return large HTML directory listings from which file links are extracted.

Strategy:
  1. GET each judgment folder (HC: Civil/Criminal/Labour/Tax; SC: Judgments)
  2. Extract .doc/.docx links via regex
  3. Download each file and extract full text with python-docx
  4. Parse parties / case number / citation / date from the filename
"""

import sys
import re
import io
import time
import logging
import urllib.parse
from pathlib import Path
from datetime import datetime, timezone
from typing import Generator, Dict, Any, List
from html import unescape

import requests
import urllib3
import docx as python_docx

# ejustice.jud.na serves a valid DigiCert / GeoTrust certificate but omits the
# intermediate, so some clients cannot build a chain. We disable verification
# for this single host; the data is public and integrity is not a concern.
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

PROJECT_ROOT = Path(__file__).resolve().parents[3]
sys.path.insert(0, str(PROJECT_ROOT))

from common.base_scraper import BaseScraper

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
)
logger = logging.getLogger("legal-data-hunter.NA.eJustice")

BASE_URL = "https://ejustice.jud.na"

FOLDERS = [
    ("High Court", "Civil",    "/High Court/Judgments/Civil/"),
    ("High Court", "Criminal", "/High Court/Judgments/Criminal/"),
    ("High Court", "Labour",   "/High Court/Judgments/Labour/"),
    ("High Court", "Tax",      "/High Court/Judgments/Tax/"),
    ("Supreme Court", None,    "/Supreme Court/Judgments/Judgments/"),
]

DOC_LINK_RE = re.compile(
    r'href="(/(?:High|Supreme) Court/Judgments/[^"]+?\.(?:docx?|DOCX?))"'
)
CASE_NUMBER_RE = re.compile(
    r"\(((?:HC|SA|CA)[A-Z\-]*[\-/ ]?\d[\d\-/. ]*[A-Z\d]*)\)",
    re.IGNORECASE,
)
CITATION_RE = re.compile(
    r"\[?(\d{4})\]?\s*(NA(?:HCMD|HCNLD|LCMD|HC|SC|CC)[A-Z]*)\s*(\d+)?",
)
DATE_RE = re.compile(
    r"\((\d{1,2})\s+(January|February|March|April|May|June|July|"
    r"August|September|October|November|December)\s+(\d{4})\)",
    re.IGNORECASE,
)
MONTHS = {
    "january": 1, "february": 2, "march": 3, "april": 4, "may": 5, "june": 6,
    "july": 7, "august": 8, "september": 9, "october": 10, "november": 11,
    "december": 12,
}


class NAeJusticeScraper(BaseScraper):

    def __init__(self):
        super().__init__(Path(__file__).parent)
        self.session = requests.Session()
        self.session.verify = False
        self.session.headers.update({
            "User-Agent": "LegalDataHunter/1.0 (Open Data Research)",
            "Accept": "text/html,application/xhtml+xml,*/*",
        })

    def _list_folder(self, folder_path: str) -> List[str]:
        url = BASE_URL + urllib.parse.quote(folder_path)
        time.sleep(1.0)
        resp = self.session.get(url, timeout=60)
        if resp.status_code != 200:
            logger.warning(f"Folder {folder_path}: HTTP {resp.status_code}")
            return []
        # The same href appears multiple times (name link, type icon, etc.).
        links = sorted(set(DOC_LINK_RE.findall(resp.text)))
        logger.info(f"Folder {folder_path}: {len(links)} unique files")
        return links

    def _download_docx_text(self, url_path: str) -> str:
        url = BASE_URL + urllib.parse.quote(url_path)
        for attempt in range(3):
            try:
                time.sleep(1.0)
                resp = self.session.get(url, timeout=120)
                if resp.status_code != 200:
                    logger.warning(f"  HTTP {resp.status_code}: {url_path}")
                    return ""
                if len(resp.content) > 50_000_000:
                    logger.warning(f"  Too large ({len(resp.content)} B): {url_path}")
                    return ""
                # python-docx only handles .docx; .doc is the OLE format.
                if not url_path.lower().endswith(".docx"):
                    logger.info(f"  Legacy .doc skipped: {url_path}")
                    return ""
                doc = python_docx.Document(io.BytesIO(resp.content))
                parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
                for t in doc.tables:
                    for row in t.rows:
                        for cell in row.cells:
                            txt = cell.text.strip()
                            if txt:
                                parts.append(txt)
                text = "\n".join(parts)
                text = re.sub(r"\n{3,}", "\n\n", text)
                return text.strip()
            except Exception as e:
                logger.warning(f"  attempt {attempt + 1}: {e}")
                if attempt < 2:
                    time.sleep(3)
        return ""

    @staticmethod
    def _parse_filename(filename: str) -> Dict[str, str]:
        m = CASE_NUMBER_RE.search(filename)
        case_number = m.group(1) if m else ""
        m = CITATION_RE.search(filename)
        citation = " ".join(g for g in m.groups() if g) if m else ""
        m = DATE_RE.search(filename)
        if m:
            day = int(m.group(1))
            month = MONTHS[m.group(2).lower()]
            year = int(m.group(3))
            date_iso = f"{year:04d}-{month:02d}-{day:02d}"
        else:
            date_iso = None
        title = filename.rsplit(".", 1)[0]
        title = re.sub(r"\s+\(\d{1,2}\s+\w+\s+\d{4}\)\s*$", "", title).strip()
        return {
            "case_number": case_number,
            "citation": citation,
            "date": date_iso,
            "title": title,
        }

    def normalize(self, raw: Dict[str, Any]) -> Dict[str, Any]:
        now = datetime.now(timezone.utc).isoformat()
        url_path = raw["url_path"]
        filename = url_path.rsplit("/", 1)[-1]
        parsed = self._parse_filename(filename)
        slug = re.sub(r"[^A-Za-z0-9]+", "-", filename.rsplit(".", 1)[0]).strip("-")
        return {
            "_id": f"NA/eJustice/{slug}",
            "_source": "NA/eJustice",
            "_type": "case_law",
            "_fetched_at": now,
            "title": parsed["title"],
            "text": raw.get("_full_text", ""),
            "date": parsed["date"],
            "url": BASE_URL + urllib.parse.quote(url_path),
            "court": raw["court"],
            "category": raw.get("category"),
            "case_number": parsed["case_number"],
            "citation": parsed["citation"],
        }

    def fetch_all(self, sample: bool = False) -> Generator[Dict[str, Any], None, None]:
        limit = 15 if sample else None
        count = 0
        seen = set()

        for court, category, folder in FOLDERS:
            if limit and count >= limit:
                break
            links = self._list_folder(folder)
            for url_path in links:
                if limit and count >= limit:
                    break
                if url_path in seen:
                    continue
                seen.add(url_path)
                text = self._download_docx_text(url_path)
                if len(text) < 500:
                    logger.warning(f"  Too short ({len(text)} chars): {url_path}")
                    continue
                count += 1
                logger.info(
                    f"  [{count}] {url_path.rsplit('/', 1)[-1][:70]} "
                    f"({len(text)} chars)"
                )
                yield {
                    "url_path": url_path,
                    "court": court,
                    "category": category,
                    "_full_text": text,
                }

        logger.info(f"Fetched {count} documents total")

    def fetch_updates(self, since: str) -> Generator[Dict[str, Any], None, None]:
        # The eJustice folder pages list newest files first; without per-file
        # modification metadata we just re-walk and let the dedup layer skip
        # records already ingested.
        yield from self.fetch_all(sample=False)


if __name__ == "__main__":
    scraper = NAeJusticeScraper()

    if len(sys.argv) < 2:
        print("Usage: python bootstrap.py [bootstrap|update|test] [--sample]")
        sys.exit(1)

    command = sys.argv[1]
    sample_mode = "--sample" in sys.argv

    if command == "test":
        scraper.test_connection()
    elif command == "bootstrap":
        scraper.bootstrap(sample_mode=sample_mode)
    elif command == "update":
        scraper.bootstrap(sample_mode=False)
    else:
        print(f"Unknown command: {command}")
        sys.exit(1)
