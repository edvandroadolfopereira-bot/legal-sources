#!/usr/bin/env python3
"""
INTL/CAN-Gacetas - Gaceta Oficial del Acuerdo de Cartagena (Andean Community)

Fetches the full-text supranational legislation (normativa andina) of the
Comunidad Andina (CAN): Decisiones de la Comisión / Consejo Andino, Resoluciones
de la Secretaría General, Dictámenes, and Tratados y Protocolos. These instruments
are published in the Gaceta Oficial del Acuerdo de Cartagena and are binding on the
member states Bolivia (BO), Colombia (CO), Ecuador (EC) and Peru (PE).

Data access
-----------
The public site comunidadandina.org is a Next.js App Router app backed by a Strapi
CMS (host https://normativa.comunidadandina.org). The Strapi REST API itself is
token-protected (403), but the Next.js server exposes a public, token-injecting
proxy route at:

    https://www.comunidadandina.org/api/normativa/?page=N&per_page=100&group=<slug>

which returns the Strapi `normative-documents` collection wrapped as
{"data": {"data": [...], "meta": {"pagination": {...}}}}. Each record carries a
`file` media object whose `url` (e.g. /uploads/DECISION_966.docx) is served from
https://normativa.comunidadandina.org. Document bodies are .docx (majority), with
some legacy .doc and a few .pdf. Full text is extracted from the downloaded file.

Groups (slugs) with content, totals as of 2026-06:
    decisiones (982), resoluciones (2601), dictamenes (149),
    tratados-y-protocolos (14)

License: Official publication of the Comunidad Andina (open government data).
"""

import argparse
import hashlib
import io
import json
import logging
import os
import re
import subprocess
import sys
import tempfile
import time
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, Iterator, Optional

import requests
import urllib3

urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger(__name__)

SOURCE_ID = "INTL/CAN-Gacetas"
PROXY_URL = "https://www.comunidadandina.org/api/normativa/"
MEDIA_BASE = "https://normativa.comunidadandina.org"
SOURCE_DIR = Path(__file__).resolve().parent

# Normative-document kinds that carry binding legislation with attached full text.
GROUPS = ["decisiones", "resoluciones", "dictamenes", "tratados-y-protocolos"]

HEADERS = {
    "User-Agent": "Mozilla/5.0 (compatible; LegalDataHunter/1.0; +https://github.com/ZachLaik)",
    "Accept": "application/json",
}

REQUEST_DELAY = 1.0
MIN_TEXT_CHARS = 200


# --------------------------------------------------------------------------- #
# Text extraction
# --------------------------------------------------------------------------- #
def _extract_docx(content: bytes) -> str:
    """Extract text from a .docx. Prefers python-docx; falls back to a
    pure-stdlib zipfile/XML parse so extraction still works on hosts where
    python-docx is not installed (the VPS image lacks it — see issue #880)."""
    try:
        from docx import Document
    except ImportError:
        return _extract_docx_stdlib(content)

    doc = Document(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    # include table cell text (some decisiones use tables for annexes)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    parts.append(t)
    return "\n".join(parts)


def _extract_docx_stdlib(content: bytes) -> str:
    """Pure-stdlib .docx text extraction (no python-docx dependency).

    A .docx is a ZIP whose word/document.xml holds the body; each <w:p> is a
    paragraph and each <w:t> a text run. Matches python-docx output closely.
    """
    import zipfile
    from xml.etree import ElementTree as ET

    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    with zipfile.ZipFile(io.BytesIO(content)) as z:
        xml = z.read("word/document.xml").decode("utf-8", "ignore")
    root = ET.fromstring(xml)
    parts = []
    for p in root.iter(ns + "p"):
        line = "".join(t.text for t in p.iter(ns + "t") if t.text).strip()
        if line:
            parts.append(line)
    return "\n".join(parts)


def _extract_pdf(content: bytes) -> str:
    import pypdf

    reader = pypdf.PdfReader(io.BytesIO(content))
    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            continue
    return "\n".join(pages)


def _extract_doc(content: bytes) -> str:
    """Best-effort extraction of legacy binary .doc files.

    Tries, in order: antiword (Linux/VPS), macOS `textutil`, then a pure-python
    fallback that pulls readable runs out of the OLE WordDocument stream. Returns
    "" if nothing usable is found.
    """
    with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tf:
        tf.write(content)
        path = tf.name
    try:
        # antiword (common on Linux servers)
        try:
            out = subprocess.run(
                ["antiword", path], capture_output=True, timeout=60
            )
            if out.returncode == 0 and out.stdout.strip():
                return out.stdout.decode("utf-8", "ignore")
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass
        # macOS textutil
        try:
            out = subprocess.run(
                ["textutil", "-convert", "txt", "-stdout", path],
                capture_output=True,
                timeout=60,
            )
            if out.returncode == 0 and out.stdout.strip():
                return out.stdout.decode("utf-8", "ignore")
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass
        # pure-python OLE fallback (needs olefile)
        try:
            ole_text = _extract_doc_ole(content)
            if ole_text.strip():
                return ole_text
        except Exception:
            pass
        # last-resort: raw-bytes recovery, no external package required
        try:
            return _extract_doc_rawbytes(content)
        except Exception:
            return ""
    finally:
        try:
            os.unlink(path)
        except OSError:
            pass


def _extract_doc_ole(content: bytes) -> str:
    """Crude pure-python .doc text recovery via olefile WordDocument stream."""
    import olefile

    if not olefile.isOleFile(io.BytesIO(content)):
        return ""
    ole = olefile.OleFileIO(io.BytesIO(content))
    if not ole.exists("WordDocument"):
        return ""
    raw = ole.openstream("WordDocument").read()
    # Decode as latin-1, keep printable + common whitespace, collapse noise.
    text = raw.decode("latin-1", "ignore")
    # keep sequences of readable characters (letters, digits, punctuation, spaces)
    chunks = re.findall(r"[A-Za-zÀ-ÿ0-9\s\.,;:%°\-\(\)/\"'¿?¡!N°]{6,}", text)
    cleaned = "\n".join(c.strip() for c in chunks if len(c.strip()) > 5)
    return cleaned


def _extract_doc_rawbytes(content: bytes) -> str:
    """Last-resort .doc recovery with no external dependency (no olefile).

    Decodes the whole OLE container as latin-1 and keeps runs of readable
    characters, then drops chunks that are mostly non-ASCII-letter noise (the
    binary streams that surround the WordDocument text). Density filtering keeps
    the real Spanish prose and discards the FF/0x00 padding runs.
    """
    text = content.decode("latin-1", "ignore")
    chunks = re.findall(r"[A-Za-zÀ-ÿ0-9\s\.,;:%°\-\(\)/\"'¿?¡!N°]{6,}", text)
    out = []
    for c in chunks:
        c = c.strip()
        if len(c) < 6:
            continue
        ascii_letters = sum(ch.isalpha() and ch.isascii() for ch in c)
        # real text is mostly latin letters; binary noise is not
        if ascii_letters < len(c) * 0.5:
            continue
        out.append(c)
    return "\n".join(out)


def extract_text(content: bytes, ext: str) -> str:
    ext = (ext or "").lower()
    try:
        if ext == ".docx":
            text = _extract_docx(content)
        elif ext == ".pdf":
            text = _extract_pdf(content)
        elif ext == ".doc":
            text = _extract_doc(content)
        else:
            return ""
    except Exception as exc:
        logger.warning("  extraction failed (%s): %s", ext, exc)
        return ""
    return _clean_text(text)


def _clean_text(text: str) -> str:
    if not text:
        return ""
    # normalize whitespace, strip stray control chars
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# --------------------------------------------------------------------------- #
# Fetcher
# --------------------------------------------------------------------------- #
class CANGacetasFetcher:
    def __init__(self):
        self.session = requests.Session()
        self.session.headers.update(HEADERS)

    def _get_page(self, group: str, page: int, per_page: int) -> Dict[str, Any]:
        for attempt in range(4):
            try:
                resp = self.session.get(
                    PROXY_URL,
                    params={"page": page, "per_page": per_page, "group": group},
                    timeout=60,
                    verify=False,
                )
                resp.raise_for_status()
                return resp.json().get("data", {}) or {}
            except Exception as exc:
                wait = 2 ** attempt
                logger.warning("  page fetch failed (%s p%d, try %d): %s; retry in %ds",
                               group, page, attempt + 1, exc, wait)
                time.sleep(wait)
        return {}

    def _download(self, url: str) -> Optional[bytes]:
        for attempt in range(3):
            try:
                resp = self.session.get(url, timeout=90, verify=False)
                resp.raise_for_status()
                return resp.content
            except Exception as exc:
                wait = 2 ** attempt
                logger.warning("  download failed (try %d): %s; retry in %ds",
                               attempt + 1, exc, wait)
                time.sleep(wait)
        return None

    def iter_raw(self, sample: bool = False) -> Iterator[Dict[str, Any]]:
        per_page = 50
        for group in GROUPS:
            page = 1
            seen = 0
            while True:
                data = self._get_page(group, page, per_page)
                records = data.get("data") or []
                if not records:
                    break
                for rec in records:
                    rec["_group"] = group
                    yield rec
                    seen += 1
                pagination = (data.get("meta") or {}).get("pagination") or {}
                page_count = pagination.get("pageCount") or 1
                logger.info("  %s page %d/%s (%d records so far)",
                            group, page, page_count, seen)
                if sample and seen >= 6:
                    break
                if page >= page_count:
                    break
                page += 1
                time.sleep(REQUEST_DELAY)
            if sample and seen >= 6:
                # gather a few from the next group too for diversity, then stop overall
                continue

    def normalize(self, raw: Dict[str, Any]) -> Optional[Dict[str, Any]]:
        file_obj = raw.get("file") or {}
        file_url = file_obj.get("url")
        ext = file_obj.get("ext")
        if not file_url:
            return None

        full_url = file_url if file_url.startswith("http") else MEDIA_BASE + file_url
        content = self._download(full_url)
        if not content:
            return None
        time.sleep(REQUEST_DELAY)

        text = extract_text(content, ext)
        if not text or len(text) < MIN_TEXT_CHARS:
            logger.info("  skipping (insufficient text %d chars): %s",
                        len(text), raw.get("nomenclature"))
            return None

        nomenclature = (raw.get("nomenclature") or "").strip()
        title_body = (raw.get("name") or "").strip()
        title = f"{nomenclature} — {title_body}" if title_body else nomenclature

        kinds = raw.get("normative_documents_kinds") or []
        kind_name = kinds[0]["name"] if kinds else raw.get("_group")

        gaceta = raw.get("gaceta") or {}
        gaceta_nom = gaceta.get("nomenclature") if isinstance(gaceta, dict) else None

        pub_date = raw.get("publication_date")  # already ISO yyyy-mm-dd or null

        strapi_id = raw.get("id") or raw.get("documentId")
        _id = f"CAN-{raw.get('_group')}-{strapi_id}"

        return {
            "_id": _id,
            "_source": SOURCE_ID,
            "_type": "legislation",
            "_fetched_at": datetime.now(timezone.utc).isoformat(),
            "title": title,
            "text": text,
            "date": pub_date,
            "url": full_url,
            "nomenclature": nomenclature,
            "document_kind": kind_name,
            "gaceta": gaceta_nom,
            "file_format": ext,
            "language": "es",
            "jurisdiction": "INTL/CAN",
            "member_states": ["BO", "CO", "EC", "PE"],
        }

    # ---- public API ---- #
    def fetch_all(self, sample: bool = False) -> Iterator[Dict[str, Any]]:
        for raw in self.iter_raw(sample=sample):
            rec = self.normalize(raw)
            if rec:
                yield rec

    def fetch_updates(self, since=None) -> Iterator[Dict[str, Any]]:
        # API returns newest first; re-fetch and let downstream dedup by _id.
        yield from self.fetch_all(sample=False)


# --------------------------------------------------------------------------- #
# CLI
# --------------------------------------------------------------------------- #
def _run_sample(fetcher: CANGacetasFetcher, target: int = 15) -> int:
    sample_dir = SOURCE_DIR / "sample"
    sample_dir.mkdir(parents=True, exist_ok=True)
    count = 0
    for rec in fetcher.fetch_all(sample=True):
        out = sample_dir / f"{rec['_id'].replace('/', '_')}.json"
        with open(out, "w", encoding="utf-8") as f:
            json.dump(rec, f, ensure_ascii=False, indent=2)
        count += 1
        logger.info("Saved %s (%d chars)", out.name, len(rec["text"]))
        if count >= target:
            break
    logger.info("Sample complete: %d records saved to %s", count, sample_dir)
    return count


def _run_full(fetcher: CANGacetasFetcher) -> int:
    data_dir = SOURCE_DIR / "data"
    data_dir.mkdir(parents=True, exist_ok=True)
    jsonl_path = data_dir / "records.jsonl"
    count = 0
    with open(jsonl_path, "w", encoding="utf-8") as jf:
        for rec in fetcher.fetch_all(sample=False):
            jf.write(json.dumps(rec, ensure_ascii=False) + "\n")
            count += 1
            if count % 50 == 0:
                logger.info("  written %d records", count)
                jf.flush()
    logger.info("Bootstrap complete: %d records written to %s", count, jsonl_path)
    print("BOOTSTRAP_COMPLETE")
    return count


def main():
    parser = argparse.ArgumentParser(description="INTL/CAN-Gacetas Fetcher")
    parser.add_argument("command", choices=["bootstrap", "bootstrap-fast"])
    parser.add_argument("--sample", action="store_true", help="Sample mode (15 records)")
    parser.add_argument("--full", action="store_true", help="Full bootstrap (JSONL)")
    args = parser.parse_args()

    fetcher = CANGacetasFetcher()
    if args.full and args.command != "bootstrap-fast":
        count = _run_full(fetcher)
    else:
        count = _run_sample(fetcher, target=15)

    if count == 0:
        logger.error("No records fetched!")
        sys.exit(1)
    sys.exit(0)


if __name__ == "__main__":
    main()
